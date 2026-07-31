"""
PDF Converter Service.
Converts PDF files to Word and Excel with advanced features:
- Page range selection
- Table extraction
- OCR support for scanned documents
- Multi-page table consolidation
- Conversion templates for recurring documents
"""

import io
import os
import tempfile
from dataclasses import dataclass, field
from typing import Any, Optional

# Fremdbibliotheken ausschließlich über app/pdf_backend.py — dort stehen
# Import, Verfügbarkeitserkennung und Lizenzregister an einer Stelle.
from app.pdf_backend import PYMUPDF_AVAILABLE, fitz, open_pdf

# Table extraction
try:
    import tabula

    TABULA_AVAILABLE = True
except ImportError:
    TABULA_AVAILABLE = False

try:
    import pandas as pd

    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

# OCR — Routing via ``app.services.ocr_router`` (ai-router oder pytesseract,
# je nach ``settings.OCR_BACKEND``). PIL bleibt hier verfuegbar fuer die
# Page-Rendering-Helper.
try:
    import pytesseract  # noqa: F401
    from PIL import Image  # noqa: F401

    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# Excel export
from app.pdf_backend import OPENPYXL_AVAILABLE

if OPENPYXL_AVAILABLE:
    import openpyxl


@dataclass
class PageInfo:
    """Information about a PDF page."""

    page_number: int
    width: float
    height: float
    has_text: bool
    has_images: bool
    has_tables: bool
    thumbnail_base64: str | None = None


@dataclass
class TableData:
    """Extracted table data from PDF."""

    page_number: int
    table_index: int
    headers: list[str]
    rows: list[list[Any]]
    confidence: float = 1.0


@dataclass
class ConversionResult:
    """Result of a PDF conversion."""

    success: bool
    output_format: str  # 'docx', 'xlsx', 'preview'
    file_content: bytes | None = None
    tables: list[TableData] = field(default_factory=list)
    preview_data: dict | None = None
    warnings: list[str] = field(default_factory=list)
    error: str | None = None


@dataclass
class ExtractionTemplate:
    """Template for recurring PDF extraction patterns."""

    id: str
    name: str
    description: str | None = None
    # Page selection pattern
    page_pattern: str | None = None  # e.g., "first", "last", "3-5"
    # Table extraction settings
    table_area: dict | None = None  # top, left, width, height in %
    skip_rows: int = 0
    skip_columns: list[int] = field(default_factory=list)
    header_row: int = 0
    # Column mapping
    column_mapping: dict[str, str] = field(default_factory=dict)


class PdfConverterService:
    """
    Service for converting PDF files to Word/Excel formats.
    Supports page selection, table extraction, and OCR.
    """

    def __init__(self):
        self.check_dependencies()

    def check_dependencies(self) -> dict[str, bool]:
        """Check which dependencies are available."""
        return {
            "pymupdf": PYMUPDF_AVAILABLE,
            "tabula": TABULA_AVAILABLE,
            "pandas": PANDAS_AVAILABLE,
            "ocr": OCR_AVAILABLE,
            "openpyxl": OPENPYXL_AVAILABLE,
        }

    def get_pdf_info(self, file_content: bytes) -> dict[str, Any]:
        """
        Get information about a PDF file.

        Returns:
            Dict with page count, pages info, and metadata
        """
        if not PYMUPDF_AVAILABLE:
            raise RuntimeError("PyMuPDF nicht installiert")

        doc = fitz.open(stream=file_content, filetype="pdf")

        pages = []
        for i, page in enumerate(doc):
            text = page.get_text()
            images = page.get_images()

            page_info = PageInfo(
                page_number=i + 1,
                width=page.rect.width,
                height=page.rect.height,
                has_text=len(text.strip()) > 0,
                has_images=len(images) > 0,
                has_tables=self._detect_tables_on_page(page),
            )
            pages.append(page_info)

        metadata = doc.metadata or {}
        doc.close()

        return {
            "page_count": len(pages),
            "pages": [vars(p) for p in pages],
            "metadata": metadata,
            "file_size": len(file_content),
        }

    def get_page_thumbnails(
        self,
        file_content: bytes,
        pages: list[int] | None = None,
        max_width: int = 200,
    ) -> dict[int, str]:
        """
        Generate thumbnail images for PDF pages.

        Args:
            file_content: PDF file bytes
            pages: List of page numbers (1-indexed), or None for all
            max_width: Maximum thumbnail width in pixels

        Returns:
            Dict mapping page number to base64-encoded PNG
        """
        if not PYMUPDF_AVAILABLE:
            raise RuntimeError("PyMuPDF nicht installiert")

        import base64

        doc = fitz.open(stream=file_content, filetype="pdf")
        thumbnails = {}

        page_nums = pages if pages else range(1, len(doc) + 1)

        for page_num in page_nums:
            if page_num < 1 or page_num > len(doc):
                continue

            page = doc[page_num - 1]

            # Calculate zoom for thumbnail
            zoom = max_width / page.rect.width
            mat = fitz.Matrix(zoom, zoom)

            # Render page to image
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")

            thumbnails[page_num] = base64.b64encode(img_bytes).decode("utf-8")

        doc.close()
        return thumbnails

    def parse_page_range(self, page_spec: str, total_pages: int) -> list[int]:
        """
        Parse a page specification string into a list of page numbers.

        Supports:
        - Single pages: "5"
        - Ranges: "3-7"
        - Comma-separated: "1,3,5-7,10"
        - Keywords: "first", "last", "all"

        Args:
            page_spec: Page specification string
            total_pages: Total number of pages in document

        Returns:
            List of page numbers (1-indexed)
        """
        page_spec = page_spec.strip().lower()

        if page_spec == "all" or page_spec == "*":
            return list(range(1, total_pages + 1))
        if page_spec == "first":
            return [1]
        if page_spec == "last":
            return [total_pages]

        pages = set()
        parts = page_spec.replace(" ", "").split(",")

        for part in parts:
            if "-" in part:
                try:
                    start, end = part.split("-", 1)
                    start = int(start) if start else 1
                    end = int(end) if end else total_pages
                    for p in range(start, end + 1):
                        if 1 <= p <= total_pages:
                            pages.add(p)
                except ValueError:
                    pass
            else:
                try:
                    p = int(part)
                    if 1 <= p <= total_pages:
                        pages.add(p)
                except ValueError:
                    pass

        return sorted(pages)

    def convert_to_word(
        self,
        file_content: bytes,
        pages: list[int] | None = None,
        use_ocr: bool = False,
        ocr_language: str = "deu",
    ) -> ConversionResult:
        """
        Convert PDF to Word document.

        Args:
            file_content: PDF file bytes
            pages: List of page numbers to convert (1-indexed)
            use_ocr: Whether to use OCR for scanned documents
            ocr_language: OCR language code ('deu' for German)

        Returns:
            ConversionResult with DOCX file bytes
        """
        warnings = []

        if not PYMUPDF_AVAILABLE:
            return ConversionResult(
                success=False, output_format="docx", error="PyMuPDF nicht installiert"
            )

        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            total_pages = len(doc)

            # Determine pages to convert
            if pages is None:
                pages = list(range(1, total_pages + 1))

            # Extract text from selected pages
            extracted_text = []
            for page_num in pages:
                if page_num < 1 or page_num > total_pages:
                    continue

                page = doc[page_num - 1]
                text = page.get_text("text")

                # If page has no text and OCR is enabled
                if not text.strip() and use_ocr and OCR_AVAILABLE:
                    text = self._ocr_page(page, ocr_language)
                    if text:
                        warnings.append(f"Seite {page_num}: OCR verwendet")

                if text.strip():
                    extracted_text.append(f"--- Seite {page_num} ---\n\n{text}")
                else:
                    warnings.append(f"Seite {page_num}: Kein Text gefunden")

            doc.close()

            if not extracted_text:
                return ConversionResult(
                    success=False,
                    output_format="docx",
                    error="Kein Text in den ausgewählten Seiten gefunden",
                )

            # Convert to DOCX using python-docx
            from docx import Document
            from docx.shared import Pt

            docx_doc = Document()
            for text_block in extracted_text:
                lines = text_block.split("\n")
                for line in lines:
                    if line.startswith("--- Seite"):
                        # Page header
                        p = docx_doc.add_paragraph()
                        run = p.add_run(line)
                        run.bold = True
                        run.font.size = Pt(14)
                    elif line.strip():
                        docx_doc.add_paragraph(line)

            # Save to bytes
            output = io.BytesIO()
            docx_doc.save(output)
            output.seek(0)

            return ConversionResult(
                success=True,
                output_format="docx",
                file_content=output.read(),
                warnings=warnings,
            )

        except Exception as e:
            return ConversionResult(
                success=False,
                output_format="docx",
                error=f"Konvertierung fehlgeschlagen: {str(e)}",
            )

    def extract_tables(
        self,
        file_content: bytes,
        pages: list[int] | None = None,
        consolidate: bool = True,
        table_area: dict | None = None,
    ) -> ConversionResult:
        """
        Extract tables from PDF and convert to Excel.

        Args:
            file_content: PDF file bytes
            pages: List of page numbers to extract from (1-indexed)
            consolidate: Whether to merge tables from multiple pages
            table_area: Optional area to extract from (top, left, width, height in %)

        Returns:
            ConversionResult with XLSX file bytes and extracted table data
        """
        warnings = []

        if not TABULA_AVAILABLE or not PANDAS_AVAILABLE:
            return ConversionResult(
                success=False,
                output_format="xlsx",
                error="Tabula oder Pandas nicht installiert",
            )

        try:
            # Write PDF to temp file for tabula
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name

            try:
                # Get total pages if not specified
                if pages is None:
                    doc = fitz.open(tmp_path)
                    pages = list(range(1, len(doc) + 1))
                    doc.close()

                # Extract tables from each page
                all_tables = []
                table_data_list = []

                for page_num in pages:
                    try:
                        # Configure extraction area if specified
                        area = None
                        if table_area:
                            # Convert percentage to points (assuming letter size)
                            area = [
                                table_area.get("top", 0) * 7.92,  # % to points
                                table_area.get("left", 0) * 6.12,
                                (
                                    table_area.get("top", 0)
                                    + table_area.get("height", 100)
                                )
                                * 7.92,
                                (
                                    table_area.get("left", 0)
                                    + table_area.get("width", 100)
                                )
                                * 6.12,
                            ]

                        # Extract tables
                        tables = tabula.read_pdf(
                            tmp_path,
                            pages=page_num,
                            area=area,
                            multiple_tables=True,
                            pandas_options={"header": None},
                        )

                        for i, df in enumerate(tables):
                            if df.empty:
                                continue

                            # Clean the dataframe
                            df = df.dropna(how="all").reset_index(drop=True)

                            if not df.empty:
                                all_tables.append(df)

                                # Create TableData
                                headers = [str(col) for col in df.columns.tolist()]
                                rows = df.values.tolist()

                                table_data_list.append(
                                    TableData(
                                        page_number=page_num,
                                        table_index=i,
                                        headers=headers,
                                        rows=rows,
                                    )
                                )

                    except Exception as e:
                        warnings.append(f"Seite {page_num}: {str(e)}")

                if not all_tables:
                    return ConversionResult(
                        success=False,
                        output_format="xlsx",
                        tables=table_data_list,
                        error="Keine Tabellen in den ausgewählten Seiten gefunden",
                        warnings=warnings,
                    )

                # Consolidate tables if requested
                if consolidate and len(all_tables) > 1:
                    # Try to merge tables with similar structure
                    consolidated = self._consolidate_tables(all_tables)
                    if consolidated is not None:
                        all_tables = [consolidated]
                        warnings.append(
                            f"{len(table_data_list)} Tabellen zusammengeführt"
                        )

                # Export to Excel
                output = io.BytesIO()
                with pd.ExcelWriter(output, engine="openpyxl") as writer:
                    for i, df in enumerate(all_tables):
                        sheet_name = (
                            f"Tabelle_{i + 1}" if len(all_tables) > 1 else "Daten"
                        )
                        df.to_excel(writer, sheet_name=sheet_name, index=False)

                output.seek(0)

                return ConversionResult(
                    success=True,
                    output_format="xlsx",
                    file_content=output.read(),
                    tables=table_data_list,
                    warnings=warnings,
                )

            finally:
                os.unlink(tmp_path)

        except Exception as e:
            return ConversionResult(
                success=False,
                output_format="xlsx",
                error=f"Tabellenextraktion fehlgeschlagen: {str(e)}",
            )

    def get_table_preview(
        self, file_content: bytes, pages: list[int] | None = None, max_rows: int = 10
    ) -> ConversionResult:
        """
        Get a preview of tables in the PDF.

        Args:
            file_content: PDF file bytes
            pages: List of page numbers to preview (1-indexed)
            max_rows: Maximum rows per table in preview

        Returns:
            ConversionResult with preview_data containing table previews
        """
        warnings = []

        if not TABULA_AVAILABLE or not PANDAS_AVAILABLE:
            return ConversionResult(
                success=False,
                output_format="preview",
                error="Tabula oder Pandas nicht installiert",
            )

        try:
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name

            try:
                if pages is None:
                    doc = fitz.open(tmp_path)
                    pages = list(range(1, len(doc) + 1))
                    doc.close()

                table_previews = []

                for page_num in pages:
                    try:
                        tables = tabula.read_pdf(
                            tmp_path,
                            pages=page_num,
                            multiple_tables=True,
                            pandas_options={"header": None},
                        )

                        for i, df in enumerate(tables):
                            if df.empty:
                                continue

                            df = df.dropna(how="all").reset_index(drop=True)

                            if not df.empty:
                                # Limit rows for preview
                                preview_df = df.head(max_rows)

                                table_previews.append(
                                    {
                                        "page": page_num,
                                        "table_index": i,
                                        "total_rows": len(df),
                                        "total_columns": len(df.columns),
                                        "headers": [
                                            str(col)
                                            for col in preview_df.columns.tolist()
                                        ],
                                        "preview_rows": preview_df.fillna(
                                            ""
                                        ).values.tolist(),
                                        "has_more": len(df) > max_rows,
                                    }
                                )

                    except Exception as e:
                        warnings.append(f"Seite {page_num}: {str(e)}")

                return ConversionResult(
                    success=True,
                    output_format="preview",
                    preview_data={
                        "tables": table_previews,
                        "total_tables": len(table_previews),
                        "pages_processed": pages,
                    },
                    warnings=warnings,
                )

            finally:
                os.unlink(tmp_path)

        except Exception as e:
            return ConversionResult(
                success=False,
                output_format="preview",
                error=f"Vorschau fehlgeschlagen: {str(e)}",
            )

    def _detect_tables_on_page(self, page) -> bool:
        """Simple heuristic to detect if page might have tables."""
        # Look for patterns that suggest tables (multiple columns of aligned text)
        dict_blocks = page.get_text("dict")
        blocks = dict_blocks.get("blocks", [])

        # Count text blocks and their positions
        x_positions = []
        for block in blocks:
            if block.get("type") == 0:  # text block
                x_positions.append(block.get("bbox", [0])[0])

        # If there are multiple distinct x positions, might be a table
        if len({round(x / 50) for x in x_positions}) > 2:
            return True

        return False

    def _ocr_page(self, page, language: str = "deu") -> str:
        """Perform OCR on a page via local Tesseract."""
        if not OCR_AVAILABLE:
            return ""

        try:
            import io as _io

            # Render page to image at 300 DPI
            mat = fitz.Matrix(300 / 72, 300 / 72)
            pix = page.get_pixmap(matrix=mat)
            img = Image.open(_io.BytesIO(pix.tobytes("png")))
            text = pytesseract.image_to_string(img, lang=language)
            return text or ""

        except Exception:
            return ""

    def _consolidate_tables(self, tables: list) -> Optional:
        """Try to consolidate multiple tables with similar structure."""
        if not tables:
            return None

        # Get first table as reference
        ref_cols = len(tables[0].columns)

        # Check if all tables have same column count
        compatible = all(len(t.columns) == ref_cols for t in tables)

        if compatible:
            # Concatenate all tables
            try:
                consolidated = pd.concat(tables, ignore_index=True)
                # Remove duplicate headers that might have been picked up
                consolidated = consolidated.drop_duplicates()
                return consolidated
            except Exception:
                return None

        return None


# Convenience function
def get_pdf_converter() -> PdfConverterService:
    """Get a PDF converter service instance."""
    return PdfConverterService()
