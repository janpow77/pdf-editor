"""
PDF & Document Tools Service.
Provides PDF manipulation (merge, split, rotate, compress, watermark, etc.)
and Word document tools (merge, diff, convert to PDF).
All operations are stateless file transformations using PyMuPDF, python-docx, Pillow.
"""

import base64
import io
import json
import logging
import subprocess
import tempfile
import threading
import zipfile
from dataclasses import dataclass, field
from datetime import UTC
from difflib import SequenceMatcher
from functools import lru_cache
from pathlib import Path

logger = logging.getLogger(__name__)

try:
    import fitz  # PyMuPDF

    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from PIL import Image

    PILLOW_AVAILABLE = True
except ImportError:
    PILLOW_AVAILABLE = False

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl

    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    import pikepdf

    PIKEPDF_AVAILABLE = True
except ImportError:
    PIKEPDF_AVAILABLE = False


@dataclass
class PdfToolResult:
    """Result from a PDF/Document tool operation."""

    success: bool
    output_format: str  # 'pdf', 'zip', 'docx', 'json', 'png', 'jpg'
    file_content: bytes | None = None
    metadata: dict | None = None
    warnings: list[str] = field(default_factory=list)
    error: str | None = None


def _parse_page_range(page_spec: str, total_pages: int) -> list[int]:
    """Parse page specification string into list of 1-indexed page numbers."""
    page_spec = page_spec.strip().lower()

    if page_spec in ("all", "*"):
        return list(range(1, total_pages + 1))
    if page_spec == "first":
        return [1]
    if page_spec == "last":
        return [total_pages]

    pages = set()
    parts = page_spec.replace(" ", "").split(",")

    for part in parts:
        if "-" in part:
            try:
                start_s, end_s = part.split("-", 1)
                start = int(start_s) if start_s else 1
                end = int(end_s) if end_s else total_pages
                for p in range(start, end + 1):
                    if 1 <= p <= total_pages:
                        pages.add(p)
            except ValueError:
                pass
        else:
            try:
                p = int(part)
                if 1 <= p <= total_pages:
                    pages.add(p)
            except ValueError:
                pass

    return sorted(pages)


def _to_pdf_date(dt) -> str:
    """Convert datetime to PDF date string format D:YYYYMMDDHHmmSS+00'00'."""
    return dt.strftime("D:%Y%m%d%H%M%S+00'00'")


_lo_available: bool | None = None
_lo_lock = threading.Lock()


def _libreoffice_available() -> bool:
    """Einmalige LibreOffice-Erkennung.

    ``libreoffice --version`` startet einen Subprozess und kostet ~90 ms. Das
    lief vorher bei jedem ``/api/health`` — also bei jedem Seitenaufruf. Die
    Installation ändert sich zur Laufzeit des Containers nicht, deshalb wird
    das Ergebnis prozessweit einmal ermittelt und gemerkt.
    """
    global _lo_available
    if _lo_available is not None:
        return _lo_available
    with _lo_lock:
        if _lo_available is None:
            try:
                result = subprocess.run(
                    ["libreoffice", "--version"],
                    capture_output=True,
                    timeout=5,
                )
                _lo_available = result.returncode == 0
            except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
                _lo_available = False
    return _lo_available


class PdfToolsService:
    """Service for PDF and document manipulation tools."""

    def check_features(self) -> dict[str, bool]:
        """Return dict of available features."""
        lo_available = _libreoffice_available()

        return {
            "pymupdf": PYMUPDF_AVAILABLE,
            "pillow": PILLOW_AVAILABLE,
            "python_docx": DOCX_AVAILABLE,
            "libreoffice": lo_available,
            "pdf_merge": PYMUPDF_AVAILABLE,
            "pdf_split": PYMUPDF_AVAILABLE,
            "pdf_rotate": PYMUPDF_AVAILABLE,
            "pdf_to_images": PYMUPDF_AVAILABLE,
            "pdf_compress": PYMUPDF_AVAILABLE,
            "pdf_watermark": PYMUPDF_AVAILABLE,
            "images_to_pdf": PYMUPDF_AVAILABLE and PILLOW_AVAILABLE,
            "pdf_page_ops": PYMUPDF_AVAILABLE,
            "pdf_metadata": PYMUPDF_AVAILABLE,
            "pdf_metadata_dates": PIKEPDF_AVAILABLE,
            "word_to_pdf": lo_available,
            "word_merge": DOCX_AVAILABLE,
            "word_diff": DOCX_AVAILABLE,
            "word_metadata": DOCX_AVAILABLE,
            "excel_metadata": OPENPYXL_AVAILABLE,
            "pikepdf": PIKEPDF_AVAILABLE,
            "pdf_protect": PIKEPDF_AVAILABLE,
            "pdf_unlock": PIKEPDF_AVAILABLE,
        }

    def get_thumbnails(
        self,
        file_content: bytes,
        pages: list[int] | None = None,
        max_width: int = 200,
    ) -> dict[str, str]:
        """Generate base64 PNG thumbnails for PDF pages."""
        if not PYMUPDF_AVAILABLE:
            return {}

        doc = fitz.open(stream=file_content, filetype="pdf")
        total = doc.page_count

        if pages is None:
            pages = list(range(1, total + 1))

        thumbnails = {}
        zoom = max_width / 612.0  # 612 = standard US Letter width in points
        mat = fitz.Matrix(zoom, zoom)

        for page_num in pages:
            if 1 <= page_num <= total:
                page = doc[page_num - 1]
                pix = page.get_pixmap(matrix=mat, alpha=False)
                thumbnails[str(page_num)] = base64.b64encode(
                    pix.tobytes("png")
                ).decode()

        doc.close()
        return thumbnails

    # ── PDF Merge ──────────────────────────────────────────────

    def merge_pdfs(
        self,
        files: list[tuple[str, bytes]],
        page_selections: dict[str, str] | None = None,
    ) -> PdfToolResult:
        """
        Merge multiple PDFs into one.

        Args:
            files: List of (filename, content) tuples
            page_selections: Optional dict mapping filename to page spec (e.g. "1-3,5")
        """
        if not PYMUPDF_AVAILABLE:
            return PdfToolResult(
                success=False, output_format="pdf", error="PyMuPDF nicht verfügbar"
            )

        if len(files) < 2:
            return PdfToolResult(
                success=False,
                output_format="pdf",
                error="Mindestens 2 PDFs erforderlich",
            )

        try:
            merged = fitz.open()
            warnings = []

            for filename, content in files:
                try:
                    src = fitz.open(stream=content, filetype="pdf")
                except Exception:
                    warnings.append(f"Konnte '{filename}' nicht öffnen, übersprungen")
                    continue

                pages_to_include = None
                if page_selections and filename in page_selections:
                    pages_to_include = _parse_page_range(
                        page_selections[filename], src.page_count
                    )

                if pages_to_include:
                    for p in pages_to_include:
                        merged.insert_pdf(src, from_page=p - 1, to_page=p - 1)
                else:
                    merged.insert_pdf(src)

                src.close()

            if merged.page_count == 0:
                return PdfToolResult(
                    success=False,
                    output_format="pdf",
                    error="Keine Seiten zum Zusammenführen",
                )

            output = merged.tobytes(deflate=True, garbage=4)
            merged.close()

            return PdfToolResult(
                success=True,
                output_format="pdf",
                file_content=output,
                warnings=warnings,
                metadata={"page_count": len(output)},
            )
        except Exception as e:
            return PdfToolResult(
                success=False,
                output_format="pdf",
                error=f"Fehler beim Zusammenführen: {e}",
            )

    # ── PDF Split ──────────────────────────────────────────────

    def split_pdf(
        self,
        file_content: bytes,
        mode: str = "ranges",
        ranges: str | None = None,
        every_n: int = 1,
    ) -> PdfToolResult:
        """
        Split a PDF into parts.

        Args:
            mode: "ranges" (custom page ranges) or "every_n" (split every N pages)
            ranges: Comma-separated ranges like "1-3,4-6,7-10" (for mode=ranges)
            every_n: Split every N pages (for mode=every_n)
        """
        if not PYMUPDF_AVAILABLE:
            return PdfToolResult(
                success=False, output_format="zip", error="PyMuPDF nicht verfügbar"
            )

        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            total = doc.page_count

            # Determine page groups
            groups: list[list[int]] = []

            if mode == "every_n":
                every_n = max(1, every_n)
                for start in range(0, total, every_n):
                    end = min(start + every_n, total)
                    groups.append(list(range(start + 1, end + 1)))
            else:
                # Parse ranges - each range becomes a separate file
                if not ranges:
                    return PdfToolResult(
                        success=False,
                        output_format="zip",
                        error="Seitenbereiche erforderlich",
                    )
                for range_spec in ranges.split(";"):
                    range_spec = range_spec.strip()
                    if range_spec:
                        pages = _parse_page_range(range_spec, total)
                        if pages:
                            groups.append(pages)

            if not groups:
                return PdfToolResult(
                    success=False,
                    output_format="zip",
                    error="Keine gültigen Seitenbereiche",
                )

            # Single group → return PDF directly
            if len(groups) == 1:
                out_doc = fitz.open()
                for p in groups[0]:
                    out_doc.insert_pdf(doc, from_page=p - 1, to_page=p - 1)
                output = out_doc.tobytes(deflate=True, garbage=4)
                out_doc.close()
                doc.close()
                return PdfToolResult(
                    success=True, output_format="pdf", file_content=output
                )

            # Multiple groups → ZIP
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                for i, group in enumerate(groups, 1):
                    out_doc = fitz.open()
                    for p in group:
                        out_doc.insert_pdf(doc, from_page=p - 1, to_page=p - 1)
                    zf.writestr(
                        f"teil_{i:03d}.pdf", out_doc.tobytes(deflate=True, garbage=4)
                    )
                    out_doc.close()

            doc.close()

            return PdfToolResult(
                success=True,
                output_format="zip",
                file_content=zip_buffer.getvalue(),
                metadata={"parts": len(groups)},
            )
        except Exception as e:
            return PdfToolResult(
                success=False, output_format="zip", error=f"Fehler beim Teilen: {e}"
            )

    # ── PDF Rotate ─────────────────────────────────────────────

    def rotate_pages(
        self,
        file_content: bytes,
        rotations: dict[int, int],
    ) -> PdfToolResult:
        """
        Rotate specific pages.

        Args:
            rotations: Dict mapping page number (1-indexed) to rotation angle (90, 180, 270)
        """
        if not PYMUPDF_AVAILABLE:
            return PdfToolResult(
                success=False, output_format="pdf", error="PyMuPDF nicht verfügbar"
            )

        try:
            doc = fitz.open(stream=file_content, filetype="pdf")

            for page_num, angle in rotations.items():
                if 1 <= page_num <= doc.page_count:
                    page = doc[page_num - 1]
                    page.set_rotation((page.rotation + angle) % 360)

            output = doc.tobytes(deflate=True, garbage=4)
            doc.close()

            return PdfToolResult(
                success=True,
                output_format="pdf",
                file_content=output,
                metadata={"rotated_pages": len(rotations)},
            )
        except Exception as e:
            return PdfToolResult(
                success=False, output_format="pdf", error=f"Fehler beim Rotieren: {e}"
            )

    # ── PDF to Images ──────────────────────────────────────────

    def pdf_to_images(
        self,
        file_content: bytes,
        pages: str | None = None,
        fmt: str = "png",
        dpi: int = 150,
    ) -> PdfToolResult:
        """Convert PDF pages to images, returned as ZIP."""
        if not PYMUPDF_AVAILABLE:
            return PdfToolResult(
                success=False, output_format="zip", error="PyMuPDF nicht verfügbar"
            )

        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            total = doc.page_count

            page_list = (
                _parse_page_range(pages, total) if pages else list(range(1, total + 1))
            )

            zoom = dpi / 72.0
            mat = fitz.Matrix(zoom, zoom)

            fmt = fmt.lower()
            if fmt not in ("png", "jpg", "jpeg"):
                fmt = "png"
            pix_fmt = "jpeg" if fmt in ("jpg", "jpeg") else "png"
            ext = "jpg" if pix_fmt == "jpeg" else "png"

            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                for page_num in page_list:
                    if 1 <= page_num <= total:
                        page = doc[page_num - 1]
                        pix = page.get_pixmap(matrix=mat, alpha=False)
                        zf.writestr(f"seite_{page_num:04d}.{ext}", pix.tobytes(pix_fmt))

            doc.close()

            return PdfToolResult(
                success=True,
                output_format="zip",
                file_content=zip_buffer.getvalue(),
                metadata={"pages": len(page_list), "format": ext, "dpi": dpi},
            )
        except Exception as e:
            return PdfToolResult(
                success=False,
                output_format="zip",
                error=f"Fehler bei Bildkonvertierung: {e}",
            )

    # ── PDF Compress ───────────────────────────────────────────

    def compress_pdf(
        self,
        file_content: bytes,
        quality: str = "medium",
    ) -> PdfToolResult:
        """
        Compress a PDF.

        Args:
            quality: "low" (max compression), "medium", "high" (minimal compression)
        """
        if not PYMUPDF_AVAILABLE:
            return PdfToolResult(
                success=False, output_format="pdf", error="PyMuPDF nicht verfügbar"
            )

        try:
            original_size = len(file_content)
            doc = fitz.open(stream=file_content, filetype="pdf")

            # Image quality settings
            image_quality_map = {"low": 40, "medium": 65, "high": 85}
            img_quality = image_quality_map.get(quality, 65)

            # Re-compress images on each page
            for page in doc:
                image_list = page.get_images(full=True)
                for img_info in image_list:
                    xref = img_info[0]
                    try:
                        base_image = doc.extract_image(xref)
                        if base_image and base_image.get("image"):
                            img_bytes = base_image["image"]
                            pil_img = Image.open(io.BytesIO(img_bytes))

                            # Convert to RGB if needed
                            if pil_img.mode in ("RGBA", "P"):
                                pil_img = pil_img.convert("RGB")

                            # Re-encode as JPEG with target quality
                            buf = io.BytesIO()
                            pil_img.save(
                                buf, format="JPEG", quality=img_quality, optimize=True
                            )
                            new_bytes = buf.getvalue()

                            # Only replace if actually smaller
                            if len(new_bytes) < len(img_bytes):
                                # Create a small PDF with the new image and insert it
                                img_rect = fitz.Rect(
                                    0, 0, pil_img.width, pil_img.height
                                )
                                img_doc = fitz.open()
                                img_page = img_doc.new_page(
                                    width=pil_img.width, height=pil_img.height
                                )
                                img_page.insert_image(img_rect, stream=new_bytes)
                                # Update the xref with compressed data
                                doc.xref_set_key(xref, "Filter", "/DCTDecode")
                                doc.xref_set_key(xref, "ColorSpace", "/DeviceRGB")
                                doc._updateStream(xref, new_bytes)
                                img_doc.close()
                    except Exception:
                        continue  # Skip images that can't be recompressed

            output = doc.tobytes(deflate=True, garbage=4, clean=True)
            doc.close()

            compressed_size = len(output)
            savings = max(0, original_size - compressed_size)
            ratio = (
                (1 - compressed_size / original_size) * 100 if original_size > 0 else 0
            )

            return PdfToolResult(
                success=True,
                output_format="pdf",
                file_content=output,
                metadata={
                    "original_size": original_size,
                    "compressed_size": compressed_size,
                    "savings_bytes": savings,
                    "compression_ratio": round(ratio, 1),
                },
            )
        except Exception as e:
            return PdfToolResult(
                success=False,
                output_format="pdf",
                error=f"Fehler beim Komprimieren: {e}",
            )

    # ── Watermark ──────────────────────────────────────────────

    def add_watermark(
        self,
        file_content: bytes,
        text: str | None = None,
        image_content: bytes | None = None,
        position: str = "center",
        opacity: float = 0.3,
        rotation: float = -45,
        font_size: float = 60,
        color: str = "#808080",
        pages: str | None = None,
    ) -> PdfToolResult:
        """Add text or image watermark to PDF pages."""
        if not PYMUPDF_AVAILABLE:
            return PdfToolResult(
                success=False, output_format="pdf", error="PyMuPDF nicht verfügbar"
            )

        if not text and not image_content:
            return PdfToolResult(
                success=False, output_format="pdf", error="Text oder Bild erforderlich"
            )

        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            total = doc.page_count

            page_list = (
                _parse_page_range(pages, total) if pages else list(range(1, total + 1))
            )

            # Parse hex color
            r, g, b = 0.5, 0.5, 0.5
            if color.startswith("#") and len(color) == 7:
                r = int(color[1:3], 16) / 255
                g = int(color[3:5], 16) / 255
                b = int(color[5:7], 16) / 255

            import math

            for page_num in page_list:
                if 1 <= page_num <= total:
                    page = doc[page_num - 1]
                    rect = page.rect

                    if text:
                        # Calculate position
                        tw = fitz.get_text_length(
                            text, fontname="helv", fontsize=font_size
                        )
                        positions = {
                            "center": fitz.Point(
                                rect.width / 2 - tw / 2, rect.height / 2
                            ),
                            "top-left": fitz.Point(50, 80),
                            "top-right": fitz.Point(rect.width - tw - 50, 80),
                            "bottom-left": fitz.Point(50, rect.height - 50),
                            "bottom-right": fitz.Point(
                                rect.width - tw - 50, rect.height - 50
                            ),
                        }
                        pos = positions.get(position, positions["center"])

                        # Build rotation matrix around page center
                        pivot = fitz.Point(rect.width / 2, rect.height / 2)
                        mat = fitz.Matrix(1, 1).prerotate(rotation)

                        page.insert_text(
                            pos,
                            text,
                            fontname="helv",
                            fontsize=font_size,
                            color=(r, g, b),
                            overlay=True,
                            morph=(pivot, mat),
                            fill_opacity=opacity,
                            stroke_opacity=opacity,
                        )

                    elif image_content:
                        # Image watermark
                        img_rect = self._calc_watermark_rect(rect, position)
                        page.insert_image(
                            img_rect,
                            stream=image_content,
                            overlay=True,
                            alpha=int(opacity * 255),
                        )

            output = doc.tobytes(deflate=True, garbage=4)
            doc.close()

            return PdfToolResult(
                success=True,
                output_format="pdf",
                file_content=output,
                metadata={"pages_watermarked": len(page_list)},
            )
        except Exception as e:
            return PdfToolResult(
                success=False,
                output_format="pdf",
                error=f"Fehler beim Wasserzeichen: {e}",
            )

    def _calc_watermark_rect(
        self, page_rect: "fitz.Rect", position: str
    ) -> "fitz.Rect":
        """Calculate rectangle for image watermark based on position."""
        w = page_rect.width * 0.3
        h = page_rect.height * 0.3
        cx, cy = page_rect.width / 2, page_rect.height / 2

        if position == "center":
            return fitz.Rect(cx - w / 2, cy - h / 2, cx + w / 2, cy + h / 2)
        elif position == "top-left":
            return fitz.Rect(30, 30, 30 + w, 30 + h)
        elif position == "top-right":
            return fitz.Rect(page_rect.width - w - 30, 30, page_rect.width - 30, 30 + h)
        elif position == "bottom-left":
            return fitz.Rect(
                30, page_rect.height - h - 30, 30 + w, page_rect.height - 30
            )
        elif position == "bottom-right":
            return fitz.Rect(
                page_rect.width - w - 30,
                page_rect.height - h - 30,
                page_rect.width - 30,
                page_rect.height - 30,
            )
        return fitz.Rect(cx - w / 2, cy - h / 2, cx + w / 2, cy + h / 2)

    # ── Images to PDF ──────────────────────────────────────────

    def images_to_pdf(
        self,
        files: list[tuple[str, bytes]],
        page_size: str = "a4",
        orientation: str = "portrait",
    ) -> PdfToolResult:
        """Convert multiple images to a single PDF."""
        if not PYMUPDF_AVAILABLE or not PILLOW_AVAILABLE:
            return PdfToolResult(
                success=False,
                output_format="pdf",
                error="PyMuPDF und Pillow erforderlich",
            )

        try:
            # Page sizes in points (72 points = 1 inch)
            sizes = {
                "a4": (595.28, 841.89),
                "a3": (841.89, 1190.55),
                "letter": (612, 792),
                "legal": (612, 1008),
            }
            pw, ph = sizes.get(page_size.lower(), sizes["a4"])
            if orientation == "landscape":
                pw, ph = ph, pw

            doc = fitz.open()
            warnings = []

            for filename, content in files:
                try:
                    # Validate image
                    img = Image.open(io.BytesIO(content))
                    img.verify()

                    # Re-open after verify
                    img = Image.open(io.BytesIO(content))
                    if img.mode in ("RGBA", "P"):
                        img = img.convert("RGB")

                    # Create page and insert image with aspect ratio
                    page = doc.new_page(width=pw, height=ph)

                    iw, ih = img.size
                    scale = min((pw - 40) / iw, (ph - 40) / ih)
                    new_w = iw * scale
                    new_h = ih * scale
                    x = (pw - new_w) / 2
                    y = (ph - new_h) / 2

                    # Convert to bytes
                    buf = io.BytesIO()
                    img.save(buf, format="JPEG", quality=90)
                    img_bytes = buf.getvalue()

                    page.insert_image(
                        fitz.Rect(x, y, x + new_w, y + new_h),
                        stream=img_bytes,
                    )
                except Exception as ex:
                    warnings.append(f"Konnte '{filename}' nicht verarbeiten: {ex}")

            if doc.page_count == 0:
                return PdfToolResult(
                    success=False,
                    output_format="pdf",
                    error="Keine Bilder konnten verarbeitet werden",
                )

            num_pages = doc.page_count
            output = doc.tobytes(deflate=True, garbage=4)
            doc.close()

            return PdfToolResult(
                success=True,
                output_format="pdf",
                file_content=output,
                warnings=warnings,
                metadata={"page_count": num_pages},
            )
        except Exception as e:
            return PdfToolResult(
                success=False, output_format="pdf", error=f"Fehler: {e}"
            )

    # ── Page Operations (delete/reorder) ───────────────────────

    def page_operations(
        self,
        file_content: bytes,
        new_order: list[int],
    ) -> PdfToolResult:
        """
        Reorder/delete pages. new_order is a list of 1-indexed page numbers
        in the desired order. Pages not included are deleted.
        """
        if not PYMUPDF_AVAILABLE:
            return PdfToolResult(
                success=False, output_format="pdf", error="PyMuPDF nicht verfügbar"
            )

        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            total = doc.page_count

            # Validate
            valid_order = [p for p in new_order if 1 <= p <= total]
            if not valid_order:
                return PdfToolResult(
                    success=False,
                    output_format="pdf",
                    error="Keine gültigen Seitennummern",
                )

            out_doc = fitz.open()
            for p in valid_order:
                out_doc.insert_pdf(doc, from_page=p - 1, to_page=p - 1)

            output = out_doc.tobytes(deflate=True, garbage=4)
            out_doc.close()
            doc.close()

            return PdfToolResult(
                success=True,
                output_format="pdf",
                file_content=output,
                metadata={
                    "original_pages": total,
                    "new_pages": len(valid_order),
                    "deleted": total - len(set(valid_order)),
                },
            )
        except Exception as e:
            return PdfToolResult(
                success=False, output_format="pdf", error=f"Fehler: {e}"
            )

    # ── Metadata ───────────────────────────────────────────────

    def get_metadata(self, file_content: bytes) -> PdfToolResult:
        """Read PDF metadata."""
        if not PYMUPDF_AVAILABLE:
            return PdfToolResult(
                success=False, output_format="json", error="PyMuPDF nicht verfügbar"
            )

        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            meta = doc.metadata or {}
            info = {
                "title": meta.get("title", ""),
                "author": meta.get("author", ""),
                "subject": meta.get("subject", ""),
                "keywords": meta.get("keywords", ""),
                "creator": meta.get("creator", ""),
                "producer": meta.get("producer", ""),
                "creation_date": meta.get("creationDate", ""),
                "modification_date": meta.get("modDate", ""),
                "page_count": doc.page_count,
                "file_size": len(file_content),
            }
            doc.close()

            return PdfToolResult(
                success=True,
                output_format="json",
                metadata=info,
            )
        except Exception as e:
            return PdfToolResult(
                success=False, output_format="json", error=f"Fehler: {e}"
            )

    def set_metadata(
        self,
        file_content: bytes,
        title: str | None = None,
        author: str | None = None,
        subject: str | None = None,
        keywords: str | None = None,
        creation_date: str | None = None,
        modification_date: str | None = None,
        creator: str | None = None,
        producer: str | None = None,
    ) -> PdfToolResult:
        """Set PDF metadata fields including dates via pikepdf."""
        from datetime import datetime

        has_date_fields = creation_date is not None or modification_date is not None

        if has_date_fields and not PIKEPDF_AVAILABLE:
            return PdfToolResult(
                success=False,
                output_format="pdf",
                error="pikepdf nicht verfügbar (benötigt für Datumsfelder)",
            )
        if not has_date_fields and not PYMUPDF_AVAILABLE:
            return PdfToolResult(
                success=False, output_format="pdf", error="PyMuPDF nicht verfügbar"
            )

        try:
            # Use pikepdf for full metadata control (dates + text fields)
            if PIKEPDF_AVAILABLE:
                pdf = pikepdf.open(io.BytesIO(file_content))

                with pdf.open_metadata() as xmp:
                    if title is not None:
                        xmp["dc:title"] = title
                    if author is not None:
                        xmp["dc:creator"] = [author]
                    if subject is not None:
                        xmp["dc:description"] = subject
                    if keywords is not None:
                        xmp["pdf:Keywords"] = keywords
                    if creator is not None:
                        xmp["xmp:CreatorTool"] = creator
                    if producer is not None:
                        xmp["pdf:Producer"] = producer
                    if creation_date is not None:
                        dt = datetime.fromisoformat(creation_date)
                        if dt.tzinfo is None:
                            dt = dt.replace(tzinfo=UTC)
                        xmp["xmp:CreateDate"] = dt.isoformat()
                    if modification_date is not None:
                        dt = datetime.fromisoformat(modification_date)
                        if dt.tzinfo is None:
                            dt = dt.replace(tzinfo=UTC)
                        xmp["xmp:ModifyDate"] = dt.isoformat()

                # Also set the Info dict for legacy PDF readers
                info = pdf.docinfo or {}
                if title is not None:
                    info[pikepdf.Name.Title] = title
                if author is not None:
                    info[pikepdf.Name.Author] = author
                if subject is not None:
                    info[pikepdf.Name.Subject] = subject
                if keywords is not None:
                    info[pikepdf.Name.Keywords] = keywords
                if creator is not None:
                    info[pikepdf.Name.Creator] = creator
                if producer is not None:
                    info[pikepdf.Name.Producer] = producer
                if creation_date is not None:
                    dt = datetime.fromisoformat(creation_date)
                    if dt.tzinfo is None:
                        dt = dt.replace(tzinfo=UTC)
                    info[pikepdf.Name.CreationDate] = _to_pdf_date(dt)
                if modification_date is not None:
                    dt = datetime.fromisoformat(modification_date)
                    if dt.tzinfo is None:
                        dt = dt.replace(tzinfo=UTC)
                    info[pikepdf.Name.ModDate] = _to_pdf_date(dt)
                pdf.docinfo = info

                buf = io.BytesIO()
                pdf.save(buf)
                pdf.close()
                output = buf.getvalue()
            else:
                # Fallback: PyMuPDF (no date writing)
                doc = fitz.open(stream=file_content, filetype="pdf")
                meta = doc.metadata or {}
                if title is not None:
                    meta["title"] = title
                if author is not None:
                    meta["author"] = author
                if subject is not None:
                    meta["subject"] = subject
                if keywords is not None:
                    meta["keywords"] = keywords
                doc.set_metadata(meta)
                output = doc.tobytes(deflate=True, garbage=4)
                doc.close()

            updated = [
                k
                for k, v in {
                    "title": title,
                    "author": author,
                    "subject": subject,
                    "keywords": keywords,
                    "creator": creator,
                    "producer": producer,
                    "creation_date": creation_date,
                    "modification_date": modification_date,
                }.items()
                if v is not None
            ]

            return PdfToolResult(
                success=True,
                output_format="pdf",
                file_content=output,
                metadata={"updated_fields": updated},
            )
        except Exception as e:
            return PdfToolResult(
                success=False, output_format="pdf", error=f"Fehler: {e}"
            )

    # ── Word to PDF ────────────────────────────────────────────

    def word_to_pdf(
        self, file_content: bytes, filename: str = "document.docx"
    ) -> PdfToolResult:
        """Convert Word document to PDF using LibreOffice headless."""
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                input_path = Path(tmpdir) / filename
                input_path.write_bytes(file_content)

                result = subprocess.run(
                    [
                        "libreoffice",
                        "--headless",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        tmpdir,
                        str(input_path),
                    ],
                    capture_output=True,
                    timeout=120,
                )

                if result.returncode != 0:
                    stderr = result.stderr.decode(errors="replace")
                    return PdfToolResult(
                        success=False,
                        output_format="pdf",
                        error=f"LibreOffice Fehler: {stderr[:500]}",
                    )

                pdf_path = input_path.with_suffix(".pdf")
                if not pdf_path.exists():
                    return PdfToolResult(
                        success=False,
                        output_format="pdf",
                        error="PDF-Ausgabe nicht gefunden",
                    )

                return PdfToolResult(
                    success=True,
                    output_format="pdf",
                    file_content=pdf_path.read_bytes(),
                )
        except subprocess.TimeoutExpired:
            return PdfToolResult(
                success=False, output_format="pdf", error="Konvertierung Timeout (120s)"
            )
        except FileNotFoundError:
            return PdfToolResult(
                success=False,
                output_format="pdf",
                error="LibreOffice nicht installiert",
            )
        except Exception as e:
            return PdfToolResult(
                success=False, output_format="pdf", error=f"Fehler: {e}"
            )

    # ── Word Merge ─────────────────────────────────────────────

    def merge_word_docs(
        self,
        files: list[tuple[str, bytes]],
        add_page_break: bool = True,
    ) -> PdfToolResult:
        """Merge multiple Word documents into one."""
        if not DOCX_AVAILABLE:
            return PdfToolResult(
                success=False, output_format="docx", error="python-docx nicht verfügbar"
            )

        if len(files) < 2:
            return PdfToolResult(
                success=False,
                output_format="docx",
                error="Mindestens 2 Dokumente erforderlich",
            )

        try:
            from docx.oxml.ns import qn

            # Start with the first document
            base_doc = DocxDocument(io.BytesIO(files[0][1]))
            warnings = []

            for filename, content in files[1:]:
                try:
                    if add_page_break:
                        base_doc.add_page_break()

                    src_doc = DocxDocument(io.BytesIO(content))

                    for element in src_doc.element.body:
                        # Skip sectPr (section properties) from source
                        if element.tag == qn("w:sectPr"):
                            continue
                        base_doc.element.body.append(element)

                except Exception as ex:
                    warnings.append(f"Konnte '{filename}' nicht einfügen: {ex}")

            buf = io.BytesIO()
            base_doc.save(buf)

            return PdfToolResult(
                success=True,
                output_format="docx",
                file_content=buf.getvalue(),
                warnings=warnings,
                metadata={"documents_merged": len(files)},
            )
        except Exception as e:
            return PdfToolResult(
                success=False, output_format="docx", error=f"Fehler: {e}"
            )

    # ── Word Diff ──────────────────────────────────────────────

    def diff_word_docs(
        self,
        file_a_content: bytes,
        file_b_content: bytes,
        file_a_name: str = "Dokument A",
        file_b_name: str = "Dokument B",
    ) -> PdfToolResult:
        """Compare two Word documents and return a diff."""
        if not DOCX_AVAILABLE:
            return PdfToolResult(
                success=False, output_format="json", error="python-docx nicht verfügbar"
            )

        try:
            doc_a = DocxDocument(io.BytesIO(file_a_content))
            doc_b = DocxDocument(io.BytesIO(file_b_content))

            # Extract paragraphs
            paras_a = [p.text for p in doc_a.paragraphs]
            paras_b = [p.text for p in doc_b.paragraphs]

            # Use SequenceMatcher for diff
            matcher = SequenceMatcher(None, paras_a, paras_b)
            changes = []
            stats = {"added": 0, "removed": 0, "changed": 0, "unchanged": 0}

            for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                if tag == "equal":
                    stats["unchanged"] += i2 - i1
                    for k in range(i1, i2):
                        changes.append(
                            {
                                "type": "equal",
                                "line_a": k + 1,
                                "line_b": j1 + (k - i1) + 1,
                                "text": paras_a[k],
                            }
                        )
                elif tag == "replace":
                    stats["changed"] += max(i2 - i1, j2 - j1)
                    for k in range(i1, i2):
                        changes.append(
                            {
                                "type": "removed",
                                "line_a": k + 1,
                                "text": paras_a[k],
                            }
                        )
                    for k in range(j1, j2):
                        changes.append(
                            {
                                "type": "added",
                                "line_b": k + 1,
                                "text": paras_b[k],
                            }
                        )
                elif tag == "delete":
                    stats["removed"] += i2 - i1
                    for k in range(i1, i2):
                        changes.append(
                            {
                                "type": "removed",
                                "line_a": k + 1,
                                "text": paras_a[k],
                            }
                        )
                elif tag == "insert":
                    stats["added"] += j2 - j1
                    for k in range(j1, j2):
                        changes.append(
                            {
                                "type": "added",
                                "line_b": k + 1,
                                "text": paras_b[k],
                            }
                        )

            similarity = matcher.ratio()

            return PdfToolResult(
                success=True,
                output_format="json",
                metadata={
                    "file_a": file_a_name,
                    "file_b": file_b_name,
                    "paragraphs_a": len(paras_a),
                    "paragraphs_b": len(paras_b),
                    "similarity": round(similarity * 100, 1),
                    "stats": stats,
                    "changes": changes,
                },
            )
        except Exception as e:
            return PdfToolResult(
                success=False, output_format="json", error=f"Fehler: {e}"
            )

    # ── Word Metadata ─────────────────────────────────────────

    def get_word_metadata(self, file_content: bytes) -> PdfToolResult:
        """Read Word document metadata (core properties)."""
        if not DOCX_AVAILABLE:
            return PdfToolResult(
                success=False,
                output_format="json",
                error="python-docx nicht verfügbar",
            )

        try:
            doc = DocxDocument(io.BytesIO(file_content))
            props = doc.core_properties

            info = {
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "keywords": props.keywords or "",
                "category": props.category or "",
                "comments": props.comments or "",
                "last_modified_by": props.last_modified_by or "",
                "revision": props.revision or 0,
                "created": props.created.isoformat() if props.created else "",
                "modified": props.modified.isoformat() if props.modified else "",
                "last_printed": (
                    props.last_printed.isoformat() if props.last_printed else ""
                ),
            }

            return PdfToolResult(
                success=True,
                output_format="json",
                metadata=info,
            )
        except Exception as e:
            return PdfToolResult(
                success=False, output_format="json", error=f"Fehler: {e}"
            )

    def set_word_metadata(
        self,
        file_content: bytes,
        title: str | None = None,
        author: str | None = None,
        subject: str | None = None,
        keywords: str | None = None,
        category: str | None = None,
        comments: str | None = None,
        last_modified_by: str | None = None,
        modified: str | None = None,
        created: str | None = None,
    ) -> PdfToolResult:
        """Set Word document metadata fields."""
        if not DOCX_AVAILABLE:
            return PdfToolResult(
                success=False,
                output_format="docx",
                error="python-docx nicht verfügbar",
            )

        try:
            from datetime import datetime

            doc = DocxDocument(io.BytesIO(file_content))
            props = doc.core_properties

            updated = []
            if title is not None:
                props.title = title
                updated.append("title")
            if author is not None:
                props.author = author
                updated.append("author")
            if subject is not None:
                props.subject = subject
                updated.append("subject")
            if keywords is not None:
                props.keywords = keywords
                updated.append("keywords")
            if category is not None:
                props.category = category
                updated.append("category")
            if comments is not None:
                props.comments = comments
                updated.append("comments")
            if last_modified_by is not None:
                props.last_modified_by = last_modified_by
                updated.append("last_modified_by")
            if modified is not None:
                props.modified = datetime.fromisoformat(modified)
                updated.append("modified")
            if created is not None:
                props.created = datetime.fromisoformat(created)
                updated.append("created")

            buf = io.BytesIO()
            doc.save(buf)

            return PdfToolResult(
                success=True,
                output_format="docx",
                file_content=buf.getvalue(),
                metadata={"updated_fields": updated},
            )
        except Exception as e:
            return PdfToolResult(
                success=False, output_format="docx", error=f"Fehler: {e}"
            )

    # ── Excel Metadata ────────────────────────────────────────

    def get_excel_metadata(self, file_content: bytes) -> PdfToolResult:
        """Read Excel document metadata (core properties)."""
        if not OPENPYXL_AVAILABLE:
            return PdfToolResult(
                success=False,
                output_format="json",
                error="openpyxl nicht verfügbar",
            )

        try:
            wb = openpyxl.load_workbook(io.BytesIO(file_content))
            props = wb.properties

            info = {
                "title": props.title or "",
                "creator": props.creator or "",
                "subject": props.subject or "",
                "keywords": props.keywords or "",
                "category": props.category or "",
                "description": props.description or "",
                "last_modified_by": props.lastModifiedBy or "",
                "revision": props.revision or "",
                "created": props.created.isoformat() if props.created else "",
                "modified": props.modified.isoformat() if props.modified else "",
                "last_printed": (
                    props.lastPrinted.isoformat() if props.lastPrinted else ""
                ),
                "sheet_count": len(wb.sheetnames),
                "sheet_names": wb.sheetnames,
            }
            wb.close()

            return PdfToolResult(
                success=True,
                output_format="json",
                metadata=info,
            )
        except Exception as e:
            return PdfToolResult(
                success=False, output_format="json", error=f"Fehler: {e}"
            )

    def set_excel_metadata(
        self,
        file_content: bytes,
        title: str | None = None,
        creator: str | None = None,
        subject: str | None = None,
        keywords: str | None = None,
        category: str | None = None,
        description: str | None = None,
        last_modified_by: str | None = None,
        modified: str | None = None,
        created: str | None = None,
    ) -> PdfToolResult:
        """Set Excel document metadata fields."""
        if not OPENPYXL_AVAILABLE:
            return PdfToolResult(
                success=False,
                output_format="xlsx",
                error="openpyxl nicht verfügbar",
            )

        try:
            from datetime import datetime

            wb = openpyxl.load_workbook(io.BytesIO(file_content))
            props = wb.properties

            updated = []
            if title is not None:
                props.title = title
                updated.append("title")
            if creator is not None:
                props.creator = creator
                updated.append("creator")
            if subject is not None:
                props.subject = subject
                updated.append("subject")
            if keywords is not None:
                props.keywords = keywords
                updated.append("keywords")
            if category is not None:
                props.category = category
                updated.append("category")
            if description is not None:
                props.description = description
                updated.append("description")
            if last_modified_by is not None:
                props.lastModifiedBy = last_modified_by
                updated.append("last_modified_by")
            if modified is not None:
                props.modified = datetime.fromisoformat(modified)
                updated.append("modified")
            if created is not None:
                props.created = datetime.fromisoformat(created)
                updated.append("created")

            buf = io.BytesIO()
            wb.save(buf)
            wb.close()

            return PdfToolResult(
                success=True,
                output_format="xlsx",
                file_content=buf.getvalue(),
                metadata={"updated_fields": updated},
            )
        except Exception as e:
            return PdfToolResult(
                success=False, output_format="xlsx", error=f"Fehler: {e}"
            )

    # === PDF Passwortschutz (pikepdf) ===

    def protect_pdf(
        self,
        file_content: bytes,
        user_password: str,
        owner_password: str | None = None,
    ) -> PdfToolResult:
        """Encrypt a PDF with a user password (AES-256, R=6)."""
        if not PIKEPDF_AVAILABLE:
            return PdfToolResult(
                success=False, output_format="pdf", error="pikepdf nicht verfügbar"
            )
        try:
            with pikepdf.open(io.BytesIO(file_content)) as pdf:
                buf = io.BytesIO()
                pdf.save(
                    buf,
                    encryption=pikepdf.Encryption(
                        user=user_password,
                        owner=owner_password or user_password,
                        R=6,
                    ),
                )
            return PdfToolResult(
                success=True, output_format="pdf", file_content=buf.getvalue()
            )
        except pikepdf.PasswordError:
            return PdfToolResult(
                success=False,
                output_format="pdf",
                error="PDF ist bereits passwortgeschützt — bitte zuerst entsperren",
            )
        except Exception as e:
            return PdfToolResult(
                success=False, output_format="pdf", error=f"Fehler: {e}"
            )

    def unlock_pdf(self, file_content: bytes, password: str) -> PdfToolResult:
        """Remove encryption from a PDF using the known password."""
        if not PIKEPDF_AVAILABLE:
            return PdfToolResult(
                success=False, output_format="pdf", error="pikepdf nicht verfügbar"
            )
        try:
            with pikepdf.open(io.BytesIO(file_content), password=password) as pdf:
                buf = io.BytesIO()
                pdf.save(buf)
            return PdfToolResult(
                success=True, output_format="pdf", file_content=buf.getvalue()
            )
        except pikepdf.PasswordError:
            return PdfToolResult(
                success=False, output_format="pdf", error="Falsches Passwort"
            )
        except Exception as e:
            return PdfToolResult(
                success=False, output_format="pdf", error=f"Fehler: {e}"
            )


@lru_cache(maxsize=1)
def get_pdf_tools() -> PdfToolsService:
    """Get singleton instance of PdfToolsService."""
    return PdfToolsService()
