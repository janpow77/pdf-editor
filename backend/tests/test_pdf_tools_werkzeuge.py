"""Tests für bisher ungetestete pdf-tools-Endpunkte.

Abgedeckt: PDF→Bilder, Bilder→PDF, Seitenoperationen (Reihenfolge/Löschen),
Rotation und Wasserzeichen im Direktaufruf sowie die Word-/Excel-Werkzeuge
(word-to-pdf, word-merge, word-diff, word-metadata, excel-metadata).
"""

import io
import json
import zipfile

import fitz
import pytest
from fastapi.testclient import TestClient

from app.main import app
from app.services.pdf_tools_service import get_pdf_tools
from tests.test_smoke import make_pdf


@pytest.fixture(scope="module")
def client():
    with TestClient(app) as c:
        yield c


def make_docx(*absaetze: str) -> bytes:
    """Word-Dokument mit den übergebenen Absätzen (python-docx)."""
    from docx import Document

    doc = Document()
    for text in absaetze:
        doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_bild(fmt: str = "PNG", groesse=(200, 100), farbe=(200, 30, 30)) -> bytes:
    """Kleines Testbild über Pillow erzeugen."""
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", groesse, farbe).save(buf, format=fmt)
    return buf.getvalue()


def _word_to_pdf_verfuegbar() -> bool:
    """True nur, wenn LibreOffice inkl. Writer wirklich DOCX→PDF wandelt —
    das Binary allein reicht nicht (Sandbox ohne libreoffice-writer)."""
    return get_pdf_tools().word_to_pdf(make_docx("Probe"), "probe.docx").success


# ── PDF → Bilder ──────────────────────────────────────────────


def test_to_images_liefert_zip_mit_gewaehlten_seiten(client):
    from PIL import Image

    pdf = make_pdf(3, "Bildquelle")
    r = client.post(
        "/api/pdf-tools/to-images",
        data={"pages": "1-2", "format": "png", "dpi": "96"},
        files={"file": ("doku.pdf", pdf, "application/pdf")},
    )
    assert r.status_code == 200, r.text
    assert "doku_bilder.zip" in r.headers["content-disposition"]
    with zipfile.ZipFile(io.BytesIO(r.content)) as zf:
        # Nur die angeforderten Seiten 1-2, Seite 3 fehlt
        assert zf.namelist() == ["seite_0001.png", "seite_0002.png"]
        img = Image.open(io.BytesIO(zf.read("seite_0001.png")))
        # A4 (595.28 x 841.89 pt) bei 96 dpi → ca. 794 x 1123 Pixel
        assert abs(img.size[0] - 794) <= 1 and abs(img.size[1] - 1123) <= 1


def test_to_images_jpg_format(client):
    r = client.post(
        "/api/pdf-tools/to-images",
        data={"pages": "1", "format": "jpg"},
        files={"file": ("d.pdf", make_pdf(1, "JPG-Test"), "application/pdf")},
    )
    assert r.status_code == 200, r.text
    with zipfile.ZipFile(io.BytesIO(r.content)) as zf:
        assert zf.namelist() == ["seite_0001.jpg"]
        assert zf.read("seite_0001.jpg")[:2] == b"\xff\xd8"  # JPEG-Magic


def test_to_images_nur_pdf_erlaubt_400(client):
    r = client.post(
        "/api/pdf-tools/to-images",
        files={"file": ("notiz.txt", b"kein pdf", "text/plain")},
    )
    assert r.status_code == 400


# ── Bilder → PDF ──────────────────────────────────────────────


def test_images_to_pdf_baut_a4_seiten(client):
    r = client.post(
        "/api/pdf-tools/images-to-pdf",
        data={"page_size": "a4", "orientation": "portrait"},
        files=[
            ("files", ("rot.png", make_bild("PNG"), "image/png")),
            ("files", ("blau.jpg", make_bild("JPEG", (100, 200), (30, 30, 200)), "image/jpeg")),
        ],
    )
    assert r.status_code == 200, r.text
    with fitz.open(stream=r.content, filetype="pdf") as doc:
        assert doc.page_count == 2
        # A4 hochkant und pro Seite genau ein eingebettetes Bild
        assert round(doc[0].rect.width) == 595 and round(doc[0].rect.height) == 842
        assert len(doc[0].get_images(full=True)) == 1
        assert len(doc[1].get_images(full=True)) == 1


def test_images_to_pdf_lehnt_fremdformat_ab(client):
    r = client.post(
        "/api/pdf-tools/images-to-pdf",
        files=[("files", ("notiz.txt", b"kein bild", "text/plain"))],
    )
    assert r.status_code == 400
    assert "Ungültiges Bildformat" in r.json()["detail"]


# ── Seitenoperationen (Reihenfolge + Löschen) ─────────────────


def test_page_operations_reihenfolge_und_loeschen(client):
    pdf = make_pdf(3, "Ordnung")
    r = client.post(
        "/api/pdf-tools/page-operations",
        data={"new_order": json.dumps([3, 1])},  # Seite 2 wird gelöscht
        files={"file": ("d.pdf", pdf, "application/pdf")},
    )
    assert r.status_code == 200, r.text
    with fitz.open(stream=r.content, filetype="pdf") as doc:
        assert doc.page_count == 2
        assert "Seite 3" in doc[0].get_text()
        assert "Seite 1" in doc[1].get_text()
        assert all("Seite 2" not in p.get_text() for p in doc)


def test_page_operations_ungueltiges_json_400(client):
    r = client.post(
        "/api/pdf-tools/page-operations",
        data={"new_order": "kaputt"},
        files={"file": ("d.pdf", make_pdf(1, "x"), "application/pdf")},
    )
    assert r.status_code == 400


# ── Rotation (Direktaufruf, nicht Batch) ──────────────────────


def test_rotate_einzelne_seite(client):
    pdf = make_pdf(2, "Rotation")
    r = client.post(
        "/api/pdf-tools/rotate",
        data={"rotations": json.dumps({"1": 90})},
        files={"file": ("d.pdf", pdf, "application/pdf")},
    )
    assert r.status_code == 200, r.text
    assert "d_rotiert.pdf" in r.headers["content-disposition"]
    with fitz.open(stream=r.content, filetype="pdf") as doc:
        assert doc[0].rotation == 90
        assert doc[1].rotation == 0  # unangeforderte Seite bleibt unberührt
        assert "Rotation Seite 1" in doc[0].get_text()


def test_rotate_ungueltiger_winkel_400(client):
    r = client.post(
        "/api/pdf-tools/rotate",
        data={"rotations": json.dumps({"1": 45})},
        files={"file": ("d.pdf", make_pdf(1, "x"), "application/pdf")},
    )
    assert r.status_code == 400


# ── Wasserzeichen (Direktaufruf, nicht Batch) ─────────────────


def test_watermark_text_nur_auf_gewaehlten_seiten(client):
    pdf = make_pdf(2, "Inhalt")
    r = client.post(
        "/api/pdf-tools/watermark",
        data={"text": "ENTWURF", "opacity": "0.4", "pages": "1"},
        files={"file": ("d.pdf", pdf, "application/pdf")},
    )
    assert r.status_code == 200, r.text
    assert "d_wasserzeichen.pdf" in r.headers["content-disposition"]
    with fitz.open(stream=r.content, filetype="pdf") as doc:
        assert "ENTWURF" in doc[0].get_text()
        assert "ENTWURF" not in doc[1].get_text()
        assert "Inhalt Seite 1" in doc[0].get_text()  # Original bleibt lesbar


def test_watermark_ohne_text_und_bild_400(client):
    r = client.post(
        "/api/pdf-tools/watermark",
        files={"file": ("d.pdf", make_pdf(1, "x"), "application/pdf")},
    )
    assert r.status_code == 400


# ── Word → PDF ────────────────────────────────────────────────


@pytest.mark.skipif(
    not _word_to_pdf_verfuegbar(), reason="LibreOffice(-Writer) nicht verfügbar"
)
def test_word_to_pdf_konvertiert_text(client):
    docx = make_docx("Konvertierungsprobe für den Prüfbericht")
    r = client.post(
        "/api/pdf-tools/word-to-pdf",
        files={"file": ("bericht.docx", docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert r.status_code == 200, r.text
    with fitz.open(stream=r.content, filetype="pdf") as doc:
        assert doc.page_count >= 1
        assert "Konvertierungsprobe" in doc[0].get_text()


def test_word_to_pdf_nur_docx_erlaubt_400(client):
    r = client.post(
        "/api/pdf-tools/word-to-pdf",
        files={"file": ("notiz.txt", b"kein word", "text/plain")},
    )
    assert r.status_code == 400


# ── Word zusammenführen ───────────────────────────────────────


def test_word_merge_haengt_dokumente_aneinander(client):
    from docx import Document

    r = client.post(
        "/api/pdf-tools/word-merge",
        data={"add_page_break": "true"},
        files=[
            ("files", ("a.docx", make_docx("Erster Absatz A"), "application/octet-stream")),
            ("files", ("b.docx", make_docx("Zweiter Absatz B"), "application/octet-stream")),
        ],
    )
    assert r.status_code == 200, r.text
    doc = Document(io.BytesIO(r.content))
    texte = [p.text for p in doc.paragraphs]
    assert "Erster Absatz A" in texte and "Zweiter Absatz B" in texte
    assert texte.index("Erster Absatz A") < texte.index("Zweiter Absatz B")


def test_word_merge_braucht_zwei_dateien(client):
    r = client.post(
        "/api/pdf-tools/word-merge",
        files=[("files", ("a.docx", make_docx("allein"), "application/octet-stream"))],
    )
    assert r.status_code == 400


# ── Word vergleichen ──────────────────────────────────────────


def test_word_diff_meldet_aenderungen(client):
    a = make_docx("Unverändert", "Alter Satz")
    b = make_docx("Unverändert", "Neuer Satz", "Zusätzlicher Satz")
    r = client.post(
        "/api/pdf-tools/word-diff",
        files={
            "file_a": ("a.docx", a, "application/octet-stream"),
            "file_b": ("b.docx", b, "application/octet-stream"),
        },
    )
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["paragraphs_a"] == 2 and body["paragraphs_b"] == 3
    assert body["stats"]["unchanged"] == 1
    assert 0 < body["similarity"] < 100
    entfernt = [c["text"] for c in body["changes"] if c["type"] == "removed"]
    hinzu = [c["text"] for c in body["changes"] if c["type"] == "added"]
    assert "Alter Satz" in entfernt
    assert "Neuer Satz" in hinzu and "Zusätzlicher Satz" in hinzu


# ── Word-Metadaten ────────────────────────────────────────────


def test_word_metadata_setzen_und_lesen(client):
    docx = make_docx("Inhalt")
    r = client.post(
        "/api/pdf-tools/word-metadata",
        data={"title": "Prüfvermerk", "author": "Erika Musterfrau", "keywords": "EFRE, Beleg"},
        files={"file": ("v.docx", docx, "application/octet-stream")},
    )
    assert r.status_code == 200, r.text
    assert "v_meta.docx" in r.headers["content-disposition"]

    r = client.post(
        "/api/pdf-tools/word-metadata/read",
        files={"file": ("v.docx", r.content, "application/octet-stream")},
    )
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["title"] == "Prüfvermerk"
    assert body["author"] == "Erika Musterfrau"
    assert body["keywords"] == "EFRE, Beleg"


# ── Excel-Metadaten ───────────────────────────────────────────


def make_xlsx() -> bytes:
    """Arbeitsmappe mit zwei Blättern und Kern-Eigenschaften (openpyxl)."""
    import openpyxl

    wb = openpyxl.Workbook()
    wb.active.title = "Daten"
    wb.active["A1"] = "Betrag"
    wb.create_sheet("Auswertung")
    wb.properties.title = "Belegliste"
    wb.properties.creator = "Erika Musterfrau"
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def test_excel_metadata_lesen(client):
    r = client.post(
        "/api/pdf-tools/excel-metadata/read",
        files={"file": ("liste.xlsx", make_xlsx(), "application/octet-stream")},
    )
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["title"] == "Belegliste"
    assert body["creator"] == "Erika Musterfrau"
    assert body["sheet_count"] == 2
    assert body["sheet_names"] == ["Daten", "Auswertung"]


def test_excel_metadata_setzen(client):
    import openpyxl

    r = client.post(
        "/api/pdf-tools/excel-metadata",
        data={"title": "Neue Belegliste", "creator": "Max Mustermann"},
        files={"file": ("liste.xlsx", make_xlsx(), "application/octet-stream")},
    )
    assert r.status_code == 200, r.text
    assert "liste_meta.xlsx" in r.headers["content-disposition"]
    wb = openpyxl.load_workbook(io.BytesIO(r.content))
    assert wb.properties.title == "Neue Belegliste"
    assert wb.properties.creator == "Max Mustermann"
    # Blattinhalt bleibt erhalten
    assert wb["Daten"]["A1"].value == "Betrag"


def test_excel_metadata_nur_xlsx_erlaubt_400(client):
    r = client.post(
        "/api/pdf-tools/excel-metadata/read",
        files={"file": ("notiz.txt", b"kein excel", "text/plain")},
    )
    assert r.status_code == 400
